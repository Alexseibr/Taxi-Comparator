#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

OUT = ".local/docs/newstat-system-T012.docx"

doc = Document()
for s in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
    try:
        doc.styles[s].font.name = "Calibri"
        doc.styles[s].font.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    except Exception:
        pass
doc.styles["Normal"].font.size = Pt(11)

for sec in doc.sections:
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.8)

def H(level, text):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x10, 0x2A, 0x43)
    return p

def P(text=""):
    return doc.add_paragraph(text)

def MONO(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    return p

def BUL(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")

def NUM(items):
    for it in items:
        doc.add_paragraph(it, style="List Number")

def TABLE(headers, rows, widths_cm=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = "" if val is None else str(val)
            for p in t.rows[ri].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if widths_cm:
        for ci, w in enumerate(widths_cm):
            for ri in range(len(t.rows)):
                t.rows[ri].cells[ci].width = Cm(w)
    return t

# ─────────────────────── ТИТУЛ ───────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Newstat — антифрод/аналитика rwbtaxi.by")
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(0x10, 0x2A, 0x43)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Техническая документация (T012)")
r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run(
    f"Версия документа: {datetime.now().strftime('%Y-%m-%d')}\n"
    "Объект: модуль /newstat в составе rwbtaxi.by (Минск)\n"
    "Состояние кода: после отката T013 (без якорной сетки и шаблонов «книжки»)"
)
r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
doc.add_paragraph()

# ─────────────────────── СОДЕРЖАНИЕ ───────────────────────
H(1, "Содержание")
toc = [
    "1. Назначение и границы системы",
    "2. Архитектура и развёртывание",
    "3. Поток данных: загрузка → ETL → риск → UI",
    "4. Модель данных (PostgreSQL)",
    "5. Модели риска и формулы",
    "6. HTTP API",
    "7. Веб-интерфейс оператора (/newstat)",
    "8. Настройки системы",
    "9. Аутентификация, роли, безопасность",
    "10. Эксплуатация: логи, бэкапы, восстановление, релизы",
    "11. Известные ограничения и порядок развития",
    "Приложение А. Пути, порты, переменные окружения",
    "Приложение Б. Журнал миграций БД",
    "Приложение В. Схема .ods/.xlsx → /upload",
]
for line in toc:
    p = doc.add_paragraph(line); p.paragraph_format.left_indent = Cm(0.5)

doc.add_page_break()

# ─────────────────────── 1. НАЗНАЧЕНИЕ ───────────────────────
H(1, "1. Назначение и границы системы")
P(
    "Newstat — это раздел сайта rwbtaxi.by по адресу /newstat, предназначенный для финансово-"
    "антифрод-аналитики работы службы такси: контроль выплат по программе гарантий смены, "
    "обнаружение водителей с признаками накрутки заказов, клиентов с подозрительной выручкой "
    "по кэшбэку и устойчивых пар «водитель–клиент» (collusion). Источник данных — суточные "
    "выгрузки из основной операционной системы такси."
)
H(2, "1.1 Что входит в систему")
BUL([
    "Веб-интерфейс оператора по адресу https://rwbtaxi.by/newstat (логин admin, чтение/запись).",
    "REST API по адресу https://rwbtaxi.by/api/newstat/* (внутренний сервис на 127.0.0.1:3012).",
    "Изолированная база данных PostgreSQL rwbtaxi_newstat (отдельная от основного сайта и парсера тарифов).",
    "ETL-задачи, которые автоматически пересчитывают ежедневные агрегаты и риск-метрики после загрузки данных.",
    "Файловые бэкапы и журналы (/var/log/rwbtaxi-newstat.log, бэкапы dist).",
])
H(2, "1.2 Что НЕ входит и НЕ должно изменяться")
BUL([
    "Парсер тарифов rwbtaxi-screens (сервис на :3011) и калибровка rwbtaxi-calib (:3010) — это отдельные системы, не зависящие от newstat.",
    "Основная (операционная) база rwbtaxi и сайт-витрина — newstat читает данные через выгрузки, не подключаясь к чужим БД.",
    "Учётные записи водителей/клиентов в основной системе — newstat хранит свой справочник drivers/clients только для целей агрегации.",
])
H(2, "1.3 Целевая аудитория документа")
BUL([
    "Технический владелец: ведёт инциденты, обновляет правила риска, релизит фронт/бэк.",
    "Оператор-аналитик: работает в /newstat, разбирает водителей/клиентов/связки.",
    "DevOps: занимается БД, бэкапами, NGINX, systemd.",
])

# ─────────────────────── 2. АРХИТЕКТУРА ───────────────────────
H(1, "2. Архитектура и развёртывание")
H(2, "2.1 Топология компонентов")
P("Все компоненты живут на одном Hetzner VPS (Ubuntu 24.04, IP 94.130.167.173). NGINX — единственная "
  "точка входа из интернета. Внутренние сервисы слушают только loopback (127.0.0.1).")
TABLE(
    ["Компонент", "Где", "Порт", "Назначение"],
    [
        ["NGINX", "VPS", "443/80", "TLS-терминирование, маршрутизация по путям"],
        ["Фронт newstat", "/var/www/rwbtaxi/dist/public", "статика", "SPA на React+Vite, отдаётся NGINX"],
        ["rwbtaxi-newstat (Node.js)", "/opt/rwbtaxi-newstat", "127.0.0.1:3012", "API /api/newstat/*"],
        ["PostgreSQL", "локально", "5432", "БД rwbtaxi_newstat"],
        ["rwbtaxi-screens (вне зоны)", "/opt/rwbtaxi-screens", "127.0.0.1:3011", "Парсер тарифов конкурентов — НЕ ТРОГАТЬ"],
        ["rwbtaxi-calib (вне зоны)", "/opt/rwbtaxi-calib", "127.0.0.1:3010", "Калибровка тарифов — НЕ ТРОГАТЬ"],
    ],
    widths_cm=[5, 5, 3, 6],
)
H(2, "2.2 Маршрутизация NGINX")
BUL([
    "/                    → SPA (общий index.html /var/www/rwbtaxi/dist/public)",
    "/newstat, /newstat/* → тот же SPA, маршрутизация роутером React (history-fallback на index.html)",
    "/api/newstat/*       → proxy_pass на 127.0.0.1:3012 (префикс /api/newstat/ срезается NGINX, до сервиса доходит /health и т.п.)",
])
H(2, "2.3 Изоляция от остального проекта")
BUL([
    "Отдельная БД rwbtaxi_newstat и отдельный пользователь newstat_user (не имеет прав на чужие БД).",
    "Бэкенд — отдельный systemd unit rwbtaxi-newstat.service.",
    "Фронт-код в монорепе изолирован пространством имён src/newstat/ — общий с основным сайтом только Vite-build и UI-kit.",
    "Ни один из сервисов newstat не общается с rwbtaxi-screens или rwbtaxi-calib и не читает их БД.",
])
H(2, "2.4 Языки и зависимости")
BUL([
    "Бэкенд: Node.js (ESM, .mjs), express, pg, pino, zod, bcrypt.",
    "Фронт: React 18, Vite 7, React Router, TanStack Query, Tailwind, shadcn/ui.",
    "БД: PostgreSQL.",
])

# ─────────────────────── 3. ПОТОК ДАННЫХ ───────────────────────
H(1, "3. Поток данных: загрузка → ETL → риск → UI")
H(2, "3.1 Логические шаги")
NUM([
    "Оператор/скрипт отправляет POST /api/newstat/upload с массивом orders (≤ 50 000 заказов за раз).",
    "Сервис валидирует заказы Zod-схемой, апсёртит справочники drivers/clients и сами orders в одной транзакции.",
    "После транзакции импорта вызывается ETL для уникальных дат (recomputeForDates).",
    "ETL пересчитывает дневные агрегаты (daily_driver_stats, daily_client_stats, daily_pair_stats, driver_shift_attendance) и риск (driver_risk_daily, client_risk_daily, pair_risk_daily).",
    "Фронт читает агрегаты через /api/newstat/daily/*. Все таблицы обновляются идемпотентно — повторный пересчёт безопасен.",
])
H(2, "3.2 Идемпотентность и повторные загрузки")
BUL([
    "Заказ распознаётся по полю order_id. Повторная загрузка того же набора → ON CONFLICT DO UPDATE по orders.",
    "Если у заказа в новой выгрузке поменялась дата — старая дата тоже попадает в список ETL, чтобы агрегаты не «зависали» на старой дате.",
    "POST /api/newstat/recompute с массивом дат — принудительный пересчёт без новой загрузки данных (используется, например, после изменения процента кэшбэка в настройках).",
])
H(2, "3.3 Ошибки ETL")
P("Импорт orders и ETL разделены: если ETL упадёт, заказы уже сохранены, но агрегаты будут устаревшими. "
  "Сервис вернёт HTTP 207 с полем etl_ok=false и ошибкой; в этом случае нужно перезапустить /recompute "
  "для пострадавших дат вручную.")

# ─────────────────────── 4. МОДЕЛЬ ДАННЫХ ───────────────────────
H(1, "4. Модель данных (PostgreSQL)")
P("База данных rwbtaxi_newstat. Все DDL описаны в /opt/rwbtaxi-newstat/migrations/001..007. "
  "Применяются скриптом lib/run-migrations.mjs от роли postgres-superuser; рантайм-роль — newstat_user "
  "(SELECT/INSERT/UPDATE/DELETE на все таблицы).")

H(2, "4.1 Справочники и сессии")
TABLE(
    ["Таблица", "Назначение", "Ключевые поля"],
    [
        ["settings", "Глобальные настройки (cashback %, пороги риска, расписание смен)", "key (PK), value jsonb"],
        ["users", "Учётные записи операторов", "id, login, password_hash (bcrypt), role"],
        ["sessions", "Активные сессии веб-кабинета", "token (PK), user_id, expires_at"],
        ["shifts", "Каталог рабочих смен с выплатами по гарантии", "id, name, start_hour, end_hour, payout_byn, weekday_mask, active"],
        ["drivers / clients", "Внутренние справочники, накапливаются при импорте", "id, имя/телефон, meta jsonb"],
    ],
    widths_cm=[4, 8, 6],
)

H(2, "4.2 Заказы и пакеты")
TABLE(
    ["Таблица", "Назначение", "Ключевые поля"],
    [
        ["orders", "Сырые заказы из выгрузок (один заказ — одна строка)",
         "order_id (PK), order_date, status, payment_type (cash|noncash), gmv, km, driver_id, client_id, "
         "arrival_minutes, trip_minutes, lat_in/lng_in, lat_out/lng_out, batch_id, raw jsonb"],
        ["upload_batches", "Журнал загрузок", "id, uploaded_at, uploaded_by, source, total_rows, inserted_rows, duplicate_rows"],
    ],
    widths_cm=[4, 6, 8],
)

H(2, "4.3 Дневные агрегаты (ETL-таблицы)")
TABLE(
    ["Таблица", "PK", "Что считаем"],
    [
        ["daily_driver_stats", "(driver_id, date)",
         "Заказы, GMV (нал/безнал), short_trip_orders (≤ short_trip_km), fast_arrival_orders (≤ fast_arrival_min), "
         "уникальные клиенты, max_orders_with_one_client, repeat_client_ratio, активные часы (битовая маска)"],
        ["daily_client_stats", "(client_id, date)",
         "Те же метрики со стороны клиента: noncash/cash, repeat_driver_ratio, cashback_earned, fast_arrival_orders, short_trip_orders"],
        ["daily_pair_stats", "(driver_id, client_id, date)",
         "Сколько заказов прошло между парой за день, безнал, GMV, короткие/быстрые"],
        ["driver_shift_attendance", "(driver_id, date, shift_id)",
         "Фактическая отработка смены: covered_hours, attendance_pct, qualified, payout_byn"],
    ],
    widths_cm=[4.5, 4.5, 9],
)

H(2, "4.4 Дневные риск-таблицы")
TABLE(
    ["Таблица", "PK", "Содержит"],
    [
        ["driver_risk_daily", "(driver_id, date)",
         "guarantee_risk, earnings_risk, collusion_risk, total_risk, *_money_byn, money_at_risk_byn, signals jsonb"],
        ["client_risk_daily", "(client_id, date)",
         "cashback_exposure, repeat_driver_dependency, suspicious_activity, total_risk, cashback_money_byn, money_at_risk_byn, total_orders, signals"],
        ["pair_risk_daily", "(driver_id, client_id, date)",
         "repeat_ratio, suspicious_ratio, cashback_dependency, total_risk, collusion_loss_risk_byn, signals"],
    ],
    widths_cm=[4.5, 4.5, 9],
)
P("Поле signals (jsonb) — это «коробка с объяснениями»: туда выгружаются исходные коэффициенты "
  "и вклады каждого сигнала. UI рисует карточку кейса прямо из signals, не делая дополнительных SQL.")

H(2, "4.5 Индексы")
P("Везде, где идёт чтение «по дате» (главный экран /newstat) — есть индексы на колонке date "
  "(idx_dds_date, idx_dcs_date, idx_dps_date, idx_dsa_date, idx_crd_date, idx_prd_date). "
  "Сортировки топов по money_at_risk_byn или collusion_loss_risk_byn опираются на отдельные "
  "композитные индексы (ix_driver_risk_date_money, idx_crd_money_at_risk, idx_prd_loss). "
  "Заказы индексированы по (driver_id, date), (client_id, date), (driver_id, client_id, date) "
  "и (status, date).")

# ─────────────────────── 5. МОДЕЛИ РИСКА ───────────────────────
H(1, "5. Модели риска и формулы")
P("Модели сознательно простые и интерпретируемые: для каждого сигнала используется линейная "
  "функция от «нормы» к «явной аномалии» (ramp(value, start, end) → [0..1]). Все коэффициенты "
  "выражены в шкале 0..100; total_risk = MAX из частных score'ов модели; money_at_risk — "
  "финансовая оценка в BYN. Все вклады попадают в signals jsonb для разбора в UI.")
P("Ключевые пороги (settings.risk_thresholds, ключ risk_thresholds):")
TABLE(
    ["Параметр", "Назначение", "Значение по умолчанию"],
    [
        ["short_trip_km", "Заказ считается «коротким», если km ≤ N", "2 км"],
        ["fast_arrival_min", "«Быстрая подача» если водитель приехал за ≤ N минут", "3 мин"],
        ["min_attendance_pct", "Минимальная посещаемость смены, чтобы выплатить гарантию", "80 %"],
        ["high_repeat_ratio", "Порог повторности (используется в эвристиках/UI)", "0.6"],
    ],
    widths_cm=[5, 8, 5],
)

H(2, "5.1 Водительский риск (driver_risk_daily)")
P("Три отдельные модели; total_risk = MAX, money_at_risk = СУММА (модели измеряют разные деньги).")
H(3, "5.1.1 guarantee_risk — формальная отработка смены")
P("Применяется ТОЛЬКО к qualified-сменам (attendance_pct ≥ min_attendance_pct). Для не-qualified "
  "счётчик = 0 (эти деньги уже не выплачиваются). Четыре сигнала по 25 баллов:")
BUL([
    "g1 = ramp(short_ratio, 0.30..0.80) × 25 — много «коротких» поездок.",
    "g2 = ramp(fast_arrival_ratio, 0.30..0.80) × 25 — много «быстрых подач».",
    "g3 = ramp(repeat_client_ratio, 0.40..0.90) × 25 — повтор одних и тех же клиентов в смену.",
    "g4 = (1 − ramp(orders_per_qualified_hour, 0.30..1.00)) × 25 — мало заказов на час смены.",
])
P("guarantee_money_byn = payout_byn × score / 100, если score ≥ 30. Иначе 0 (порог отсечения шума).")

H(3, "5.1.2 earnings_risk — накрутка / аномальные паттерны")
BUL([
    "e1 = ramp(cancel_ratio, 0.20..0.60) × 25",
    "e2 = ramp(short_ratio, 0.40..0.90) × 25",
    "e3 = ramp(cash_ratio, 0.50..1.00) × min(1, short_ratio / 0.5) × 25  — комбинация «нал + короткие»",
    "e4 = ramp(concentration_one_client, 0.40..0.80) × 25",
])
P("earnings_money_byn = total_gmv × 0.10 × score / 100, если score ≥ 30. 10 % — оценка переплаты в подозрительной зоне.")

H(3, "5.1.3 collusion_risk — зависимость от одного клиента")
BUL([
    "c1 = ramp(concentration_one_client, 0.40..0.90) × 60",
    "c2 = ramp(repeat_client_ratio, 0.50..0.95) × 40",
])
P("collusion_money_byn = noncash_top_client_estimate × score / 100, где "
  "noncash_top_client_estimate = noncash_gmv × max_orders_with_one_client / total_orders. Порог отсечения 30.")

H(2, "5.2 Клиентский риск (client_risk_daily)")
P("Три модели; total_risk = MAX. money_at_risk считается ТОЛЬКО по кэшбэк-оси "
  "(чтобы не задвоить с pair-collusion).")

H(3, "5.2.1 cashback_exposure — насколько большая часть кэшбэка может быть фейковой")
BUL([
    "ce1 = ramp(short_ratio, 0.30..0.80) × 25",
    "ce2 = ramp(fast_arrival_ratio, 0.30..0.80) × 25",
    "ce3 = ramp(noncash_ratio, 0.70..1.00) × 25 — 100 % безнала характерно для cashback-фрода.",
    "ce4 = ramp(concentration_one_driver, 0.50..1.00) × 25 — клиент ездит почти только с одним водителем.",
])
P("cashback_money_byn = cashback_earned × score / 100, при score ≥ 30. money_at_risk_byn = cashback_money_byn.")

H(3, "5.2.2 repeat_driver_dependency — клиент завязан на одного водителя")
BUL([
    "rd1 = ramp(concentration_one_driver, 0.50..1.00) × 60",
    "rd2 = ramp(repeat_driver_ratio, 0.40..0.80) × 40",
])
P("Деньги по этой оси не считаются отдельно — это территория pair-collusion (см. ниже).")

H(3, "5.2.3 suspicious_activity — общая «странность» клиента за день")
BUL([
    "sa1 = ramp(total_orders, 8..20) × 50 — десятки заказов в день.",
    "sa2 = noncash_ratio × 30 — линейно от 0 до 30.",
    "sa3 = ramp(min(short_ratio, fast_ratio), 0.30..0.70) × 20 — комбинация «короткие + быстрые».",
])

H(2, "5.3 Pair-collusion риск (pair_risk_daily)")
P("Триплет (driver_id, client_id, date). Три score'а; total_risk = MAX. money_at_risk_byn (= collusion_loss_risk_byn) — "
  "оценка переплаты по кэшбэку, осевшей у пары.")
BUL([
    "repeat_ratio = ramp(orders_count, 3..10) × 100 — 3 заказа уже подозрительно, 10+ — почти точно сговор.",
    "suspicious_ratio = clamp(noncash_ratio × 60 + ramp(short_fast_share, 0.3..0.7) × 40, 0..100).",
    "cashback_dependency = ramp(client_share_by_pair, 0.5..1.0) × 100, где client_share_by_pair = pair.noncash / client.noncash_orders.",
])
P("collusion_loss_risk_byn = noncash_gmv × cashback_pct / 100 × total_risk / 100. "
  "Если total_risk ниже — деньги считаются почти не под риском.")

H(2, "5.4 Гарантийная программа (driver_shift_attendance)")
P("Для каждой пары (driver_id, date, shift_id) ETL считает covered_hours — сколько часов смены было покрыто "
  "хотя бы одним заказом. attendance_pct = covered_hours / shift_hours × 100. "
  "qualified = attendance_pct ≥ settings.risk_thresholds.min_attendance_pct. "
  "payout_byn = shifts.payout_byn (если qualified) или 0.")

H(2, "5.5 Принципы интерпретации")
BUL([
    "Все score'ы выражены в одинаковой шкале 0..100, поэтому таблицы «Топ риска» можно сортировать сквозно.",
    "money_at_risk_byn — это потенциальная финансовая экспозиция, а не подтверждённый ущерб. Решение принимает оператор.",
    "Каждое значение signals в БД — это «след» расчёта; UI карточки кейса (планируется в T010) рендерит из signals напрямую.",
    "При смене порогов (cashback %, min_attendance_pct и т.п.) обязательно вызвать POST /recompute для затронутых дат.",
])

# ─────────────────────── 6. API ───────────────────────
H(1, "6. HTTP API")
P("Базовый префикс: https://rwbtaxi.by/api/newstat (NGINX срезает префикс, к сервису приходит /health и т.п.). "
  "Аутентификация — Bearer-токен из POST /auth/login. Роли: admin (полный доступ), antifraud (чтение + загрузки), viewer (только чтение).")

H(2, "6.1 Аутентификация и сессии")
TABLE(
    ["Метод", "Путь", "Роль", "Назначение"],
    [
        ["POST", "/auth/login", "—", "{login, password} → {token, expires_at, user}"],
        ["GET", "/auth/me", "любая", "Текущий пользователь по Bearer-токену"],
        ["POST", "/auth/logout", "любая", "Удаляет сессию, токен инвалидируется"],
        ["GET", "/health", "—", "Проверка БД и времени"],
        ["GET", "/version", "—", "Имя и версия сервиса"],
    ],
    widths_cm=[2, 5, 3, 8],
)

H(2, "6.2 Настройки и смены")
TABLE(
    ["Метод", "Путь", "Роль", "Назначение"],
    [
        ["GET", "/settings/all", "опц.", "Все настройки (доступно фронту до логина — для отрисовки)"],
        ["GET", "/settings/:key", "опц.", "Одна настройка по ключу"],
        ["PUT", "/settings/:key", "admin", "Перезапись значения (валидация Zod-схемой по ключу)"],
        ["GET", "/shifts", "опц.", "Список смен"],
        ["POST", "/shifts", "admin", "Создание смены"],
        ["PUT", "/shifts/:id", "admin", "Обновление смены"],
        ["DELETE", "/shifts/:id", "admin", "Удаление смены"],
    ],
    widths_cm=[2, 5, 3, 8],
)

H(2, "6.3 Загрузки и пересчёт")
TABLE(
    ["Метод", "Путь", "Роль", "Назначение"],
    [
        ["POST", "/upload", "admin/antifraud",
         "{source, orders[]} (≤50 000) → импортирует и запускает ETL для затронутых дат"],
        ["POST", "/recompute", "admin", "{dates[]} (≤366) → принудительный пересчёт ETL по датам"],
        ["GET", "/batches", "любая", "Журнал последних 100 загрузок"],
        ["GET", "/orders/sample", "любая", "?date=YYYY-MM-DD&limit=N (≤500) — заказы за день"],
    ],
    widths_cm=[2, 4, 3.5, 8.5],
)

H(2, "6.4 Дневные срезы для UI")
TABLE(
    ["Метод", "Путь", "Роль", "Назначение"],
    [
        ["GET", "/daily/summary", "любая", "Главный экран: 18 агрегатов за день (заказы, GMV, выплаты, риск)"],
        ["GET", "/daily/drivers", "любая", "Топ daily_driver_stats за день"],
        ["GET", "/daily/clients", "любая", "Топ daily_client_stats (по cashback_earned)"],
        ["GET", "/daily/pairs", "любая", "Топ daily_pair_stats по orders_count"],
        ["GET", "/daily/attendance", "любая", "Гарантийная отработка смен за день"],
        ["GET", "/daily/driver-risks", "admin/antifraud", "?date,&limit=200 (≤1000) — топ водителей по money_at_risk"],
        ["GET", "/daily/client-risks", "admin/antifraud", "?date,&limit=200 (≤500) — топ клиентов по money_at_risk"],
        ["GET", "/daily/pair-risks", "admin/antifraud", "?date,&limit=200 (≤500) — топ пар по collusion_loss_risk"],
    ],
    widths_cm=[2, 4.5, 3.5, 8],
)

H(2, "6.5 Ответ /daily/summary")
P("Формат ответа: { ok: true, summary: {...} }. Поля summary:")
BUL([
    "orders_total, orders_completed — счётчики по orders.",
    "gmv_total, gmv_noncash — суммы GMV (BYN).",
    "drivers_active, clients_active — уникальные DAU.",
    "guarantee_payout, qualified_count — выплаты по гарантии.",
    "cashback_total — суммарный кэшбэк по daily_client_stats.",
    "risk_money_total = guarantee + earnings + collusion (по водителям).",
    "risk_money_guarantee, risk_money_earnings, risk_money_collusion — по моделям.",
    "risky_drivers_count — водителей с total_risk ≥ 30.",
    "cashback_loss_total, risky_clients_count — клиентский риск.",
    "collusion_loss_total, risky_pairs_count — pair-риск.",
])

# ─────────────────────── 7. UI ───────────────────────
H(1, "7. Веб-интерфейс оператора (/newstat)")
P("SPA на React Router. Все страницы кроме /newstat/login требуют валидной сессии. "
  "Layout — общий: верхняя плашка с навигацией и выходом, основной контент.")
TABLE(
    ["Маршрут", "Файл", "Назначение"],
    [
        ["/newstat/login", "NewstatLoginPage.tsx", "Форма логина (admin)"],
        ["/newstat", "NewstatHomePage.tsx", "Главная: 4 сводные карточки + Топы и быстрые ссылки"],
        ["/newstat/upload", "NewstatUploadPage.tsx", "Загрузка выгрузок (.ods/.xlsx → JSON в /upload)"],
        ["/newstat/risks", "NewstatRisksPage.tsx", "Полный топ водителей по риску (фильтры, сортировка)"],
        ["/newstat/clients-risk", "NewstatClientsRiskPage.tsx", "Полный топ клиентов по риску"],
        ["/newstat/pairs-risk", "NewstatPairsRiskPage.tsx", "Полный топ связок «водитель–клиент»"],
        ["/newstat/guarantee", "NewstatGuaranteePage.tsx", "Гарантийные смены за день: кто qualified, payout"],
        ["/newstat/settings", "NewstatSettingsPage.tsx", "Настройки кэшбэка, порогов риска, смен"],
    ],
    widths_cm=[5, 5.5, 7],
)
P("Все таблицы умеют сортироваться по money_at_risk / total_risk; даты переключаются picker'ом наверху страницы. "
  "Запросы кэшируются TanStack Query, сессия хранится в auth-store (localStorage).")

# ─────────────────────── 8. НАСТРОЙКИ ───────────────────────
H(1, "8. Настройки системы")
P("Все настройки хранятся в таблице settings (ключ-значение jsonb). Изменение валидируется Zod-схемой по ключу.")
TABLE(
    ["Ключ", "Структура value", "Назначение"],
    [
        ["cashback", "{ percent_of_noncash: number 0..100 }", "Процент кэшбэка от безнала. По умолчанию 30."],
        ["risk_thresholds",
         "{ short_trip_km, fast_arrival_min, min_attendance_pct, high_repeat_ratio }",
         "Пороги короткой/быстрой поездки, минимальной отработки смены, повторности."],
        ["shifts_default", "{ shifts: any[] }", "Резервный шаблон расписания смен (необязательно)."],
    ],
    widths_cm=[3.5, 7, 7],
)
P("Любое изменение настроек, влияющее на риск (cashback, пороги), требует POST /recompute по затронутым датам, "
  "иначе таблицы /daily/*-risks будут показывать старые числа.")

# ─────────────────────── 9. БЕЗОПАСНОСТЬ ───────────────────────
H(1, "9. Аутентификация, роли, безопасность")
H(2, "9.1 Парольный логин")
BUL([
    "Пароли хранятся как bcrypt-хэш в users.password_hash.",
    "На несуществующего/выключенного юзера сервис всё равно вызывает bcrypt с фиктивным хэшем — чтобы не утекал признак «такого логина нет» через timing.",
    "При успешном логине создаётся sessions.token (случайный), время жизни — фиксированное (см. lib/auth.mjs).",
    "Для авторизованных запросов клиент шлёт заголовок Authorization: Bearer <token>.",
])
H(2, "9.2 Роли")
TABLE(
    ["Роль", "Что может"],
    [
        ["admin", "Все эндпоинты, включая запись настроек и /recompute"],
        ["antifraud", "Загрузка данных, чтение всех риск-эндпоинтов"],
        ["viewer", "Только чтение публичных эндпоинтов и /daily/* без риск-таблиц"],
    ],
    widths_cm=[4, 12],
)
H(2, "9.3 Сетевая изоляция")
BUL([
    "Бэкенд слушает только loopback (127.0.0.1:3012). Извне доступен только через NGINX.",
    "БД rwbtaxi_newstat — отдельная от основной БД сайта; пользователь newstat_user не имеет прав на чужие БД.",
])
H(2, "9.4 Логин по умолчанию")
P("Логин: admin. Пароль установлен при инициализации (хранится в users.password_hash). "
  "При компрометации — менять через UPDATE users SET password_hash = … (bcrypt-хэш).")

# ─────────────────────── 10. ЭКСПЛУАТАЦИЯ ───────────────────────
H(1, "10. Эксплуатация: логи, бэкапы, восстановление, релизы")
H(2, "10.1 systemd")
MONO(
    "# Статус сервиса\n"
    "systemctl status rwbtaxi-newstat\n"
    "# Перезапуск\n"
    "systemctl restart rwbtaxi-newstat\n"
    "# Юнит\n"
    "/etc/systemd/system/rwbtaxi-newstat.service\n"
    "# Конфиг окружения\n"
    "/etc/rwbtaxi-newstat.env  (DATABASE_URL, SESSION_SECRET, ANCHORS_PATH и др.)"
)
H(2, "10.2 Логи")
BUL([
    "Файловый лог приложения: /var/log/rwbtaxi-newstat.log (pino JSON). Содержит и HTTP-доступ, и ошибки.",
    "systemd-журнал: journalctl -u rwbtaxi-newstat -n 100 --no-pager (стартап-события, креш-репорты).",
    "NGINX-логи: /var/log/nginx/access.log и /var/log/nginx/error.log.",
])
H(2, "10.3 Бэкапы")
BUL([
    "БД — pg_dump rwbtaxi_newstat (по расписанию администратора VPS).",
    "Перед каждым релизом бэкенда — копия файлов с суффиксом *.bak.<метка> рядом с оригиналом в /opt/rwbtaxi-newstat/.",
    "Перед каждым релизом фронта — копия каталога /var/www/rwbtaxi/dist/public.<метка>.",
])
H(2, "10.4 Релиз бэкенда")
NUM([
    "Скопировать новые .mjs (и при необходимости новые миграции) в /opt/rwbtaxi-newstat.",
    "Сделать копии заменяемых файлов как .bak.<метка> рядом.",
    "Если есть новая миграция — применить под postgres-superuser: sudo -u postgres psql rwbtaxi_newstat -f migrations/NNN_xxx.sql",
    "Выдать GRANT'ы newstat_user на новые таблицы (и колонки, если они в новых таблицах).",
    "systemctl restart rwbtaxi-newstat и проверить /health, /version.",
])
H(2, "10.5 Релиз фронта")
NUM([
    "Локально: PORT=4173 BASE_PATH=/ pnpm --filter @workspace/taxi-compare run build (соберёт dist/public).",
    "Сделать копию текущего prod-дист: mv /var/www/rwbtaxi/dist/public /var/www/rwbtaxi/dist/public.bak.<метка>.",
    "Закатать новый dist на сервер (tar over ssh; rsync на VPS не установлен).",
    "Проверить, что новый bundle отдаётся: curl -sk https://rwbtaxi.by/newstat/ | grep assets/index-",
])
H(2, "10.6 Откат релиза")
NUM([
    "Бэкенд: cp файлы из *.bak.<метка> обратно, затем systemctl restart.",
    "Если в составе релиза была миграция — выполнить обратный DDL (DROP COLUMN/INDEX/TABLE с IF EXISTS).",
    "Фронт: mv /var/www/rwbtaxi/dist/public.bak.<метка> обратно в public, NGINX перезапускать не нужно (статика).",
    "Проверить /health и список миграций (SELECT id FROM schema_migrations).",
])

# ─────────────────────── 11. ОГРАНИЧЕНИЯ И РОАДМАП ───────────────────────
H(1, "11. Известные ограничения и порядок развития")
H(2, "11.1 Ограничения текущей версии")
BUL([
    "Риск-модели опираются на эмпирические пороги; обучение по реальным меткам пока не реализовано.",
    "В UI пока нет отдельной страницы «карточка кейса» — оператору приходится ориентироваться по signals jsonb через DevTools/JSON.",
    "Drill-down с топа в карточку конкретного водителя/клиента/связки реализован частично (T009-Б ещё в работе).",
    "В выгрузках допускаются разные форматы partner-источников; всё, кроме order_id/order_date/status, формально опционально.",
    "Якорная сетка Минска и расчёт «шаблонных маршрутов из книжки» (T013) был развёрнут, но затем откачен; код задачи остался в репозитории и в /opt/rwbtaxi-newstat/migrations/008_anchors.sql.",
])
H(2, "11.2 Порядок развития (предложение)")
NUM([
    "T009-Б + T010 — единая «карточка сущности»: метаданные, заказы с пометкой подозрительных, разбор signals по правилам, история действий.",
    "T011 — вынос сервиса в отдельный репозиторий/деплой-юнит, чтобы развивать независимо от основного сайта.",
    "T013 (повторно) — вернуть якорную сетку и шаблоны книжки, опираясь на «карточку кейса» и signals.",
    "T014+ — ML-слой поверх правил (CatBoost), таблицы fraud_features_*, fraud_tickets, fraud_training_labels, fraud_model_versions, retraining + объяснимость; final_risk = 0.7 × rule + 0.3 × ml.",
])

# ─────────────────────── ПРИЛОЖЕНИЯ ───────────────────────
H(1, "Приложение А. Пути, порты, переменные окружения")
TABLE(
    ["Что", "Значение"],
    [
        ["VPS", "94.130.167.173 (Hetzner, Ubuntu 24.04)"],
        ["SSH-ключ (только локально)", "/tmp/ssh-rwbtaxi/id_ed25519"],
        ["Корень бэкенда", "/opt/rwbtaxi-newstat"],
        ["Файл окружения", "/etc/rwbtaxi-newstat.env"],
        ["БД", "rwbtaxi_newstat (Postgres, локальный)"],
        ["Рантайм-роль", "newstat_user"],
        ["Порт сервиса", "127.0.0.1:3012"],
        ["Корень фронта", "/var/www/rwbtaxi/dist/public"],
        ["Лог приложения", "/var/log/rwbtaxi-newstat.log"],
        ["Юнит systemd", "rwbtaxi-newstat.service"],
        ["URL UI", "https://rwbtaxi.by/newstat"],
        ["URL API", "https://rwbtaxi.by/api/newstat/..."],
    ],
    widths_cm=[6, 11],
)

H(1, "Приложение Б. Журнал миграций БД")
TABLE(
    ["Файл", "Назначение"],
    [
        ["001_init.sql", "Справочники users/sessions/drivers/clients, settings, shifts, orders, upload_batches"],
        ["002_etl.sql", "daily_driver_stats / daily_client_stats / daily_pair_stats"],
        ["003_attendance.sql", "driver_shift_attendance + GRANT newstat_user"],
        ["004_driver_risk.sql", "driver_risk_daily + индексы + GRANT"],
        ["005_client_risk.sql", "client_risk_daily + дозамеры в daily_client_stats"],
        ["006_client_risk_total_orders.sql", "Колонка total_orders в client_risk_daily + бэкафилл из signals"],
        ["007_pair_risk.sql", "pair_risk_daily + индексы + GRANT"],
        ["008_anchors.sql (откатили)", "Якорная сетка, dist_category, daily_distance_breakdown, template_* — в проде сейчас отсутствует"],
    ],
    widths_cm=[6, 11],
)

H(1, "Приложение В. Схема .ods/.xlsx → /upload")
P("Серверный API ожидает массив объектов orders[]. Фронт-загрузчик сам приводит выгрузку к этому формату. "
  "Ниже — поля, которые сервер принимает (passthrough пропускает дополнительные).")
TABLE(
    ["Поле", "Тип", "Обязательно", "Комментарий"],
    [
        ["order_id", "string", "да", "Уникальный идентификатор заказа из источника"],
        ["order_date", "YYYY-MM-DD", "да", "Дата заказа (берётся для дневных агрегатов)"],
        ["status", "string", "да", "completed|cancelled|... — определяет включение в completed_orders"],
        ["created_at / cancelled_at", "ISO timestamp", "нет", "Используется для часовой маски активности"],
        ["payment_type", "cash|noncash", "нет", "Базовая колонка для разделения нал/безнал"],
        ["payment_type2 / car_class_create / car_class_appoint", "string", "нет", "Технические поля партнёра, сохраняются в raw"],
        ["driver_id / driver_name", "string", "нет (по факту нужен)", "Без driver_id заказ не попадёт в driver-агрегаты"],
        ["client_id / client_phone", "string", "нет (по факту нужен)", "Без client_id заказ не попадёт в client-агрегаты"],
        ["gmv", "number (BYN)", "нет", "Без GMV — заказ не считается в денежных метриках"],
        ["km", "number", "нет", "Сравнивается с short_trip_km для метки «короткий»"],
        ["arrival_minutes", "number", "нет", "Сравнивается с fast_arrival_min для метки «быстрая подача»"],
        ["trip_minutes", "number", "нет", "Информационно"],
        ["lat_in/lng_in/lat_out/lng_out", "number", "нет", "Точки посадки/высадки (используются в T013)"],
        ["is_now", "boolean", "нет", "Заказ-сейчас vs предзаказ"],
    ],
    widths_cm=[3.5, 3, 2.5, 8],
)

# ───────── footer ─────────
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run("Конец документа")
r.italic = True; r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print("written:", OUT, os.path.getsize(OUT), "bytes")
